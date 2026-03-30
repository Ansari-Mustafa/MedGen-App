"""Extract all {{placeholder}} patterns from a .docx file."""

import re
import sys
from docx import Document


def extract_placeholders(docx_path: str) -> list[str]:
    doc = Document(docx_path)
    pattern = re.compile(r"\{\{(.+?)\}\}")
    found = set()

    # Paragraphs
    for para in doc.paragraphs:
        # Check full paragraph text (handles split runs)
        for match in pattern.findall(para.text):
            found.add(match.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for match in pattern.findall(para.text):
                        found.add(match.strip())

    # Headers and footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for para in header_footer.paragraphs:
                for match in pattern.findall(para.text):
                    found.add(match.strip())

    return sorted(found)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_placeholders.py <file.docx>")
        sys.exit(1)

    path = sys.argv[1]
    placeholders = extract_placeholders(path)

    print(f"Found {len(placeholders)} placeholder(s):\n")
    for p in placeholders:
        if p.endswith("[]"):
            print(f"  {{{{{p}}}}}  [list → bullets]")
        else:
            print(f"  {{{{{p}}}}}")
