"""
Template onboarding service.

Two primary surfaces:
- analyze_reports(reports) — single Claude Opus call ("Template Architect").
  Returns structured JSON identifying the best-formatted report and which
  spans should become Jinja placeholders for docxtpl.
- transform_docx(docx_bytes, replacements) — rewrites the chosen .docx,
  replacing each `find` span with `{{ placeholder }}` while preserving run
  formatting. List fields collapse multi-paragraph bullet blocks into a
  single `{{ placeholder[] }}` paragraph.

List fields use the `[]` suffix convention already honoured by
`extras/scripts/extract_placeholders.py` and `backend/services/report_fill.py`.
"""
from __future__ import annotations

import asyncio
import io
import json
import re
from copy import deepcopy
from typing import Any

from docx import Document

from backend.config import get_settings


# ── Structure extraction ────────────────────────────────────────────────────


def extract_paragraphs(docx_bytes: bytes) -> list[dict]:
    """Flatten a .docx into an ordered list of paragraph descriptors.

    Covers body paragraphs and table cells. Skips empty paragraphs.
    """
    doc = Document(io.BytesIO(docx_bytes))
    items: list[dict] = []

    def _add(container: str, text: str, style: str | None):
        cleaned = text.strip()
        if not cleaned:
            return
        items.append(
            {
                "idx": len(items),
                "container": container,
                "text": cleaned,
                "style": style,
            }
        )

    for para in doc.paragraphs:
        _add("body", para.text, para.style.name if para.style else None)

    for table_idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _add(f"table{table_idx}", para.text, para.style.name if para.style else None)

    return items


def format_score(paragraphs: list[dict]) -> int:
    """Rough heuristic for how well-formatted a report is.

    Prefers documents with distinct headings and table content — those carry
    more structural signal to preserve during the transform step.
    """
    score = len(paragraphs)
    for p in paragraphs:
        style = (p.get("style") or "").lower()
        if "heading" in style:
            score += 5
        if p["container"].startswith("table"):
            score += 2
    return score


# ── Architect prompt ────────────────────────────────────────────────────────


_ARCHITECT_SYSTEM = """You are a Template Architect helping a UK consulting doctor convert their past medical-legal reports into a reusable DOCX template. You MUST return strict JSON only — no prose, no markdown fences."""


def _build_architect_prompt(reports: list[dict]) -> str:
    blocks = []
    for r in reports:
        lines = [f"[Report {r['index']}] filename={r['filename']} format_score={r['format_score']}"]
        for p in r["paragraphs"]:
            tag = p["container"]
            style = p.get("style") or "-"
            lines.append(f"  ({tag} · {style}) {p['text']}")
        blocks.append("\n".join(lines))
    reports_blob = "\n\n".join(blocks)

    return f"""You receive {len(reports)} past medical-legal reports from the same UK consulting doctor. Each is flattened into ordered paragraph descriptors below.

TASKS:
1. Pick ONE report to use as the template — prefer the one with the richest formatting (headings, tables, consistent styles). Return its index.
2. Identify every span that is DYNAMIC per-patient: names, dates, addresses, symptoms, clinical findings, narrative sections, opinions, etc. Static boilerplate (letterhead, disclaimers, signature blocks, statement of truth) must be kept verbatim.
3. For each dynamic span, return:
   - report_index: the report you extracted the span from (usually the chosen report; if a dynamic section is easier to spot in a different report, that's fine — the transformer only operates on the chosen report so the span MUST also appear there).
   - find: the EXACT substring from the chosen report that should be replaced. Must match character-for-character (copy the text verbatim). For list/bullet content, include the full multi-line bullet block joined with "\\n".
   - placeholder: snake_case name. LIST fields (bullet sections such as symptoms, recommendations, medications) MUST end with `[]`.
   - type: "string" | "list".
   - description: one short sentence describing what the field holds.
4. Capture the doctor's writing style: tone, recurring phrases, British English markers, structural habits.

OUTPUT CONTRACT (return JSON with these exact keys):
{{
  "chosen_report_index": <int, one of the report indices above>,
  "writing_style": {{
    "tone": "...",
    "britishisms": ["..."],
    "typical_phrases": ["..."]
  }},
  "doctor_profile": {{
    "structure": "...",
    "notes": "..."
  }},
  "replacements": [
    {{"report_index": 1, "find": "...", "placeholder": "patient_name", "type": "string", "description": "Full name of the claimant"}},
    {{"report_index": 1, "find": "• Headache\\n• Neck pain", "placeholder": "symptoms[]", "type": "list", "description": "Presenting symptoms as bullet list"}}
  ]
}}

REPORTS:
{reports_blob}
"""


def _extract_json(text: str) -> dict:
    """Parse JSON from a model response, tolerating fenced blocks."""
    if "```" in text:
        text = text[text.index("```"):]
        text = text.split("\n", 1)[1] if "\n" in text else text
        text = text.rsplit("```", 1)[0]
    text = text.strip()
    if not text.startswith("{"):
        start = text.find("{")
        end = text.rfind("}") + 1
        if start != -1 and end > start:
            text = text[start:end]
    # strict=False lets raw newlines/tabs through string values — Claude
    # often emits those when quoting multi-line `find` spans verbatim.
    return json.loads(text, strict=False)


async def analyze_reports(reports: list[dict]) -> dict:
    """Call Claude to analyze past reports and return the Architect JSON contract.

    `reports` shape: [{"index": 1, "filename": str, "paragraphs": [{...}], "format_score": int}, ...]
    """
    settings = get_settings()

    def _run() -> dict:
        from anthropic import Anthropic

        client = Anthropic(api_key=settings.anthropic_api_key)
        response = client.messages.create(
            model=settings.ai_model,
            max_tokens=16000,
            system=_ARCHITECT_SYSTEM,
            messages=[{"role": "user", "content": _build_architect_prompt(reports)}],
        )
        return _extract_json(response.content[0].text)

    return await asyncio.to_thread(_run)


# ── DOCX transform ──────────────────────────────────────────────────────────


def _normalise(s: str) -> str:
    # Collapse whitespace so fuzzy matches survive double-spaces / curly quotes etc.
    return re.sub(r"\s+", " ", s).strip()


def _replace_in_paragraph(paragraph, find: str, replacement: str) -> bool:
    """Replace `find` inside a paragraph's text while preserving the formatting
    of the first matched run. Clears trailing matched runs.

    Returns True if a replacement happened.
    """
    runs = paragraph.runs
    if not runs:
        return False

    full_text = "".join(r.text or "" for r in runs)
    idx = full_text.find(find)
    if idx == -1:
        # Try normalised match
        norm_full = _normalise(full_text)
        norm_find = _normalise(find)
        norm_idx = norm_full.find(norm_find)
        if norm_idx == -1:
            return False
        # Re-locate approximately using case/space-insensitive search
        # Fall back to cheap replace on the whole paragraph text.
        new_text = full_text.replace(full_text, full_text)
        if norm_find in norm_full:
            # Simple strategy: just wipe paragraph text and insert replacement.
            for r in runs[1:]:
                r.text = ""
            runs[0].text = re.sub(re.escape(_normalise(find)), replacement, _normalise(full_text), count=1)
            return True
        return False

    end = idx + len(find)

    # Walk runs to map character offsets.
    cursor = 0
    new_runs_text: list[str] = []
    replaced = False
    for run in runs:
        run_text = run.text or ""
        run_start = cursor
        run_end = cursor + len(run_text)

        if run_end <= idx or run_start >= end:
            # This run is entirely outside the match range.
            new_runs_text.append(run_text)
        else:
            prefix = run_text[: max(0, idx - run_start)] if run_start < idx else ""
            suffix = run_text[max(0, end - run_start):] if run_end > end else ""
            if not replaced:
                new_runs_text.append(prefix + replacement + suffix)
                replaced = True
            else:
                new_runs_text.append(prefix + suffix)
        cursor = run_end

    for run, new_text in zip(runs, new_runs_text):
        run.text = new_text
    return True


def _iter_paragraph_containers(doc):
    """Yield every paragraph collection in the document (body, tables, headers, footers)."""
    yield doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.paragraphs
    for section in doc.sections:
        yield section.header.paragraphs
        yield section.footer.paragraphs


def _replace_anywhere(doc, find: str, replacement: str) -> bool:
    """Replace the first occurrence of `find` anywhere in the document."""
    for paragraphs in _iter_paragraph_containers(doc):
        for para in paragraphs:
            if find in para.text or _normalise(find) in _normalise(para.text):
                if _replace_in_paragraph(para, find, replacement):
                    return True
    return False


def _collapse_list_match(doc, find: str, placeholder_tag: str) -> bool:
    """Handle a list replacement that spans multiple paragraphs (bullet blocks).

    Collapses the matched paragraph block into a single paragraph containing
    `placeholder_tag`. Returns True on success.
    """
    lines = [ln for ln in find.split("\n") if ln.strip()]
    if len(lines) <= 1:
        return _replace_anywhere(doc, find, placeholder_tag)

    # Walk body paragraphs looking for a contiguous run matching all lines.
    body_paragraphs = doc.paragraphs
    normalised_lines = [_normalise(ln) for ln in lines]

    for start in range(len(body_paragraphs) - len(lines) + 1):
        window = body_paragraphs[start : start + len(lines)]
        if all(_normalise(p.text).find(normalised_lines[i]) != -1 for i, p in enumerate(window)):
            # Replace first paragraph's text with the placeholder tag.
            first = window[0]
            for run in first.runs[1:]:
                run.text = ""
            if first.runs:
                first.runs[0].text = placeholder_tag
            else:
                first.add_run(placeholder_tag)
            # Clear subsequent matched paragraphs. Cannot easily remove them
            # without rewriting XML; setting text to "" is acceptable for
            # docxtpl rendering (empty paragraphs render as blank lines).
            for p in window[1:]:
                for run in p.runs:
                    run.text = ""
            return True

    # Fallback: single-paragraph replacement with joined text.
    return _replace_anywhere(doc, find, placeholder_tag)


async def transform_docx(docx_bytes: bytes, replacements: list[dict]) -> bytes:
    """Apply replacements to a .docx and return new bytes.

    `replacements` should all target the chosen report. Each entry:
        {find, placeholder, type: "string"|"list", description?}
    """

    def _run() -> bytes:
        doc = Document(io.BytesIO(docx_bytes))
        for repl in replacements:
            placeholder = repl["placeholder"].strip()
            if not placeholder:
                continue
            tag_name = placeholder
            # LIST fields must end with [] per docxtpl convention used elsewhere.
            if repl.get("type") == "list" and not tag_name.endswith("[]"):
                tag_name = f"{tag_name}[]"
            jinja_tag = "{{ " + tag_name + " }}"
            find = repl.get("find") or ""
            if not find:
                continue
            if repl.get("type") == "list":
                _collapse_list_match(doc, find, jinja_tag)
            else:
                _replace_anywhere(doc, find, jinja_tag)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    return await asyncio.to_thread(_run)


def placeholders_to_schema(replacements: list[dict]) -> dict[str, dict]:
    """Convert a list of replacements into the `{name: {type, description}}` dict
    stored on Template.placeholders."""
    schema: dict[str, dict] = {}
    for repl in replacements:
        name = repl.get("placeholder", "").strip()
        if not name:
            continue
        key = name[:-2] if name.endswith("[]") else name
        schema[key] = {
            "type": repl.get("type", "string"),
            "description": repl.get("description", ""),
        }
    return schema
